import spacy
from docx import Document
from docx.shared import RGBColor
from googletrans import Translator

def colorize_word(run, pos):
    # Generate a unique color for each POS tag
    color = hash(pos) % 0x1000000  # Use hash value as RGB color
    run.font.color.rgb = RGBColor(color & 0xFF, (color >> 8) & 0xFF, (color >> 16) & 0xFF)

def translate_sentence(sentence):
    translator = Translator()
    translation = translator.translate(sentence, src='en', dest='es')
    return translation.text

def colorize_spanish_words(paragraph, spanish_text):
    # Load the SpaCy model for Spanish
    nlp = spacy.load('es_core_news_sm')

    # Process the translated Spanish text
    spacy_doc = nlp(spanish_text)

    # Colorize words in Spanish
    for token in spacy_doc:
        run = paragraph.add_run(str(token.text) + " ")
        colorize_word(run, token.pos_)

if __name__ == "__main__":
    # Get user input
    input_text = "This is a text. He eats an apple."

    # Create a Word document
    doc = Document()

    # Process the text and add sentences to the Word document
    docx_paragraph = doc.add_paragraph()
    spacy_doc = spacy.load('en_core_web_sm')(input_text)

    for sent in spacy_doc.sents:
        # Analyze the sentence and colorize words in English
        for token in sent:
            run = docx_paragraph.add_run(str(token.text) + " ")
            colorize_word(run, token.pos_)

        # Translate the sentence to Spanish
        translated_sentence = translate_sentence(str(sent))

        # Add the translated sentence next to the original with a separation of '-'
        docx_paragraph.add_run('- ')

        # Colorize words in Spanish
        colorize_spanish_words(docx_paragraph, translated_sentence)

        # Add a line break between sentences
        docx_paragraph.add_run('\n')

    # Save the Word document
    doc.save('output.docx')

    print("Word document saved successfully.")
